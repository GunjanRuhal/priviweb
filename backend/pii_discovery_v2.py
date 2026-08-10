from pathlib import Path
from presidio_analyzer import (
    AnalyzerEngine,
    EntityRecognizer,
    Pattern,
    PatternRecognizer,
    RecognizerResult,
)
import pdfplumber
from docx import Document
import pandas as pd
import pytesseract
from pytesseract import Output
from pdf2image import convert_from_path
from PIL import Image
import re
from collections import defaultdict
from dataclasses import dataclass, field
from functools import lru_cache
from typing import Dict, List, Set, Tuple, Optional
import json
from datetime import datetime

import networkx as nx
from rapidfuzz import fuzz

try:
    import spacy
except ImportError:
    spacy = None

nlp = None
if spacy is not None:
    for model_name in ("en_core_web_trf", "en_core_web_sm"):
        try:
            nlp = spacy.load(model_name)
            break
        except Exception:
            continue

# ============================================================================
# CONFIGURATION
# ============================================================================

SCORE_THRESHOLD = 0.5
POPPLER_PATH = None

ALLOWED_ENTITIES = {
    # Custom Indian PII recognizers
    "IN_AADHAAR", "IN_PAN", "IN_PASSPORT", "IN_VOTER_ID", "IN_DRIVING_LICENSE",
    "IN_VEHICLE_REGISTRATION", "IN_GSTIN", "IN_PHONE", "BANK_ACCOUNT", "IFSC_CODE",
    "UPI_ID", "DATE_OF_BIRTH", "GENDER", "IP_ADDRESS", "MAC_ADDRESS", "MEDICAL_RECORD",
    # Date recognizers
    "CARD_EXPIRY_DATE", "DOCUMENT_ISSUE_DATE", "DOCUMENT_EXPIRY_DATE",
    # Presidio default
    "EMAIL_ADDRESS",
    "LABEL_MATCHED_NAME", "KIN_NAME",
}

# Confidence thresholds for PII clustering
CONFIDENCE_HIGH = 0.8
CONFIDENCE_MEDIUM = 0.5
CONFIDENCE_LOW = 0.3

# Fuzzy name matching thresholds (rapidfuzz.fuzz.token_sort_ratio, 0-100 scale)
FUZZY_ANCHOR_NAME_THRESHOLD = 88
FUZZY_ANCHOR_MIN_TOKENS = 2
FUZZY_ANCHOR_MIN_CHARS = 6
FUZZY_RECONCILIATION_THRESHOLD = 82

TXT_WINDOW_LINES = 3

# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class PIIInstance:
    """Represents a single PII detection"""
    entity_type: str
    value: str
    file_path: str
    context: str = ""
    score: float = 0.0
    record_id: str = ""
    relation: str = ""

    def __hash__(self):
        return hash((self.entity_type, self.value.lower().strip()))

    def __eq__(self, other):
        if not isinstance(other, PIIInstance):
            return False
        return (self.entity_type == other.entity_type and
                self.value.lower().strip() == other.value.lower().strip())

@dataclass
class PIICluster:
    """Represents a group of PII instances that likely belong to the same person"""
    pii_instances: List[PIIInstance] = field(default_factory=list)
    anchor_matches: Set[str] = field(default_factory=set)
    files: Set[str] = field(default_factory=set)
    confidence: float = 0.0

    def add_instance(self, instance: PIIInstance):
        self.pii_instances.append(instance)
        self.files.add(instance.file_path)

    def get_pii_summary(self) -> Dict[str, int]:
        """Get count of each PII type"""
        summary = defaultdict(int)
        for pii in self.pii_instances:
            summary[pii.entity_type] += 1
        return dict(summary)

    def get_unique_pii(self) -> Dict[str, Set[str]]:
        """Get unique values for each PII type"""
        unique = defaultdict(set)
        for pii in self.pii_instances:
            unique[pii.entity_type].add(pii.value)
        return dict(unique)

@dataclass
class TextRecord:
    """A single co-occurrence unit used as the record boundary for the anchor-proximity graph."""
    record_id: str
    text: str
    boundary_type: str
    locator: str = ""

# ============================================================================
# NAME-CANDIDATE SPAN HELPERS
# ============================================================================

_INTRA_SPAN_WS = r"[^\S\n]+"
_NAME_SPAN = r"[A-Z][a-zA-Z.'-]*(?:" + _INTRA_SPAN_WS + r"[A-Z][a-zA-Z.'-]*){0,4}"

_FIELD_LABEL_PATTERN = re.compile(
    r"(?:Name|Full Name|Account Holder|Patient Name|Applicant Name|"
    r"Employee Name|Policyholder|Signed by)\s*[:\-]\s*(" + _NAME_SPAN + r")"
)

def _find_label_matched_name_spans(text: str) -> List[Tuple[str, int, int]]:
    """Field-label name spans only (feeds LABEL_MATCHED_NAME)."""
    spans = []
    for m in _FIELD_LABEL_PATTERN.finditer(text):
        span_text = m.group(1)
        spans.append((span_text, m.start(1), m.end(1)))
    return spans

# Kin-name extraction: each pattern is paired with the relation it implies.
_KIN_LABEL_PATTERNS = [
    (re.compile(r"(?:Mother's Name|Mother Name)\s*[:\-]\s*(" + _NAME_SPAN + r")"), "mother"),
    (re.compile(r"(?:Father's Name|Father Name)\s*[:\-]\s*(" + _NAME_SPAN + r")"), "father"),
    (re.compile(r"Guardian Name\s*[:\-]\s*(" + _NAME_SPAN + r")"), "guardian"),
    (re.compile(r"Spouse Name\s*[:\-]\s*(" + _NAME_SPAN + r")"), "spouse"),
    (re.compile(r"Husband's Name\s*[:\-]\s*(" + _NAME_SPAN + r")"), "husband"),
    (re.compile(r"Wife's Name\s*[:\-]\s*(" + _NAME_SPAN + r")"), "wife"),
    (re.compile(r"Next of Kin\s*[:\-]\s*(" + _NAME_SPAN + r")"), "next_of_kin"),
]

_KIN_RELATIONAL_MARKER_PATTERNS = [
    (re.compile(r"S/o\s*[:\-]?\s*(" + _NAME_SPAN + r")"), "father"),
    (re.compile(r"D/o\s*[:\-]?\s*(" + _NAME_SPAN + r")"), "father"),
    (re.compile(r"W/o\s*[:\-]?\s*(" + _NAME_SPAN + r")"), "husband"),
    (re.compile(r"C/o\s*[:\-]?\s*(" + _NAME_SPAN + r")"), "guardian"),
]

def _find_kin_name_spans(text: str) -> List[Tuple[str, int, int, str]]:
    """Kin field-label and relational-marker name spans, with offsets and inferred relation."""
    spans = []
    for pattern, relation in _KIN_LABEL_PATTERNS + _KIN_RELATIONAL_MARKER_PATTERNS:
        for m in pattern.finditer(text):
            spans.append((m.group(1), m.start(1), m.end(1), relation))
    return spans

@lru_cache(maxsize=1024)
def _get_spacy_doc(text: str):
    """Single cached entry point for running spaCy on a given text."""
    if nlp is None:
        return None
    return nlp(text)

def _spacy_person_spans(text: str) -> List[Tuple[str, int, int]]:
    """spaCy PERSON entity spans, with offsets."""
    doc = _get_spacy_doc(text)
    if doc is None:
        return []
    return [
        (ent.text.strip(), ent.start_char, ent.end_char)
        for ent in doc.ents if ent.label_ == "PERSON"
    ]

def _name_candidate_spans(text: str) -> List[Tuple[str, int, int]]:
    """Combined candidate spans for fuzzy name matching: spaCy PERSON + label/relational-marker matches."""
    return _spacy_person_spans(text) + _find_label_matched_name_spans(text)

def _passes_fuzzy_guardrails(anchor_name: str, candidate: str) -> bool:
    if len(anchor_name.split()) < FUZZY_ANCHOR_MIN_TOKENS:
        return False
    if len(candidate.split()) < FUZZY_ANCHOR_MIN_TOKENS:
        return False
    if len(candidate) < FUZZY_ANCHOR_MIN_CHARS:
        return False
    return True

def name_matches_fuzzy(text: str, anchor_name: str) -> Optional[str]:
    """Return the first candidate span that fuzzy-matches anchor_name above
    threshold, or None. Used only as a fallback after an exact match fails."""
    for candidate, _, _ in _name_candidate_spans(text):
        if not _passes_fuzzy_guardrails(anchor_name, candidate):
            continue
        score = fuzz.token_sort_ratio(anchor_name.lower(), candidate.lower())
        if score >= FUZZY_ANCHOR_NAME_THRESHOLD:
            return candidate
    return None

# ============================================================================
# CUSTOM RECOGNIZERS
# ============================================================================

def create_comprehensive_recognizers():
    """Create all custom PII recognizers (India-specific and general)"""
    recognizers = []

    # INDIAN GOVERNMENT IDs

    aadhaar_recognizer = PatternRecognizer(
        supported_entity="IN_AADHAAR",
        name="Aadhaar Recognizer",
        patterns=[
            Pattern(name="aadhaar_pattern", regex=r"\b\d{4}\s?\d{4}\s?\d{4}\b", score=0.6)
        ],
        context=["aadhaar", "aadhar", "uid", "uidai"]
    )
    recognizers.append(aadhaar_recognizer)

    pan_recognizer = PatternRecognizer(
        supported_entity="IN_PAN",
        name="PAN Recognizer",
        patterns=[
            Pattern(name="pan_pattern", regex=r"\b[A-Z]{5}\d{4}[A-Z]\b", score=0.7)
        ],
        context=["pan", "permanent account number", "income tax"]
    )
    recognizers.append(pan_recognizer)

    passport_recognizer = PatternRecognizer(
        supported_entity="IN_PASSPORT",
        name="Indian Passport Recognizer",
        patterns=[
            Pattern(name="passport_pattern", regex=r"\b[A-Z]\d{7}\b", score=0.6)
        ],
        context=["passport", "passport number", "travel document"]
    )
    recognizers.append(passport_recognizer)

    voter_id_recognizer = PatternRecognizer(
        supported_entity="IN_VOTER_ID",
        name="Voter ID Recognizer",
        patterns=[
            Pattern(name="voter_id_pattern", regex=r"\b[A-Z]{3}\d{7}\b", score=0.6)
        ],
        context=["voter", "voter id", "epic", "election card"]
    )
    recognizers.append(voter_id_recognizer)

    dl_recognizer = PatternRecognizer(
        supported_entity="IN_DRIVING_LICENSE",
        name="Driving License Recognizer",
        patterns=[
            Pattern(name="dl_pattern", regex=r"\b[A-Z]{2}\d{2}\s?\d{11}\b", score=0.6)
        ],
        context=["driving license", "dl", "driver's license", "licence"]
    )
    recognizers.append(dl_recognizer)

    vehicle_recognizer = PatternRecognizer(
        supported_entity="IN_VEHICLE_REGISTRATION",
        name="Vehicle Registration Recognizer",
        patterns=[
            Pattern(name="vehicle_pattern", regex=r"\b[A-Z]{2}\s?\d{2}\s?[A-Z]{1,2}\s?\d{4}\b", score=0.6)
        ],
        context=["vehicle", "registration", "rc", "car number", "vehicle number"]
    )
    recognizers.append(vehicle_recognizer)

    gstin_recognizer = PatternRecognizer(
        supported_entity="IN_GSTIN",
        name="GSTIN Recognizer",
        patterns=[
            Pattern(name="gstin_pattern", regex=r"\b\d{2}[A-Z]{5}\d{4}[A-Z]\d[Z][A-Z\d]\b", score=0.7)
        ],
        context=["gstin", "gst", "tax", "goods and services tax"]
    )
    recognizers.append(gstin_recognizer)

    # CONTACT INFORMATION

    indian_phone_recognizer = PatternRecognizer(
        supported_entity="IN_PHONE",
        name="Indian Phone Recognizer",
        patterns=[
            Pattern(name="indian_phone_pattern", regex=r"\b[6-9]\d{9}\b", score=0.5),
            Pattern(name="indian_phone_with_code", regex=r"\+91[\s-]?[6-9]\d{9}\b", score=0.7)
        ],
        context=["phone", "mobile", "contact", "call", "whatsapp", "telephone"]
    )
    recognizers.append(indian_phone_recognizer)

    # FINANCIAL DATA

    bank_account_recognizer = PatternRecognizer(
        supported_entity="BANK_ACCOUNT",
        name="Bank Account Recognizer",
        patterns=[
            Pattern(name="bank_account_pattern", regex=r"\b\d{9,18}\b", score=0.3)
        ],
        context=["account number", "bank account", "a/c no", "account no", "acc no", "savings account", "current account"]
    )
    recognizers.append(bank_account_recognizer)

    ifsc_recognizer = PatternRecognizer(
        supported_entity="IFSC_CODE",
        name="IFSC Code Recognizer",
        patterns=[
            Pattern(name="ifsc_pattern", regex=r"\b[A-Z]{4}0[A-Z0-9]{6}\b", score=0.8)
        ],
        context=["ifsc", "bank code", "branch code", "ifsc code"]
    )
    recognizers.append(ifsc_recognizer)

    upi_recognizer = PatternRecognizer(
        supported_entity="UPI_ID",
        name="UPI ID Recognizer",
        patterns=[
            Pattern(name="upi_pattern", regex=r"\b[\w\.-]+@(?:paytm|phonepe|googlepay|gpay|ybl|okaxis|okicici|okhdfcbank|oksbi|axl|ibl|airtel)\b", score=0.7)
        ],
        context=["upi", "upi id", "payment address", "paytm", "phonepe", "gpay"]
    )
    recognizers.append(upi_recognizer)

    # PERSONAL INFORMATION

    dob_recognizer = PatternRecognizer(
        supported_entity="DATE_OF_BIRTH",
        name="Date of Birth Recognizer",
        patterns=[
            Pattern(name="dob_ddmmyyyy", regex=r"\b\d{2}[/-]\d{2}[/-]\d{4}\b", score=0.3),
            Pattern(name="dob_yyyymmdd", regex=r"\b\d{4}[/-]\d{2}[/-]\d{2}\b", score=0.3)
        ],
        context=["dob", "date of birth", "birth date", "born on", "birthday", "d.o.b", "birth", "born"]
    )
    recognizers.append(dob_recognizer)

    card_expiry_recognizer = PatternRecognizer(
        supported_entity="CARD_EXPIRY_DATE",
        name="Card Expiry Date Recognizer",
        patterns=[
            Pattern(name="card_expiry_mmyy", regex=r"\b(0[1-9]|1[0-2])/(\d{2}|\d{4})\b", score=0.3)
        ],
        context=["expiry", "exp", "valid thru", "valid till", "thru", "till"]
    )
    recognizers.append(card_expiry_recognizer)

    issue_date_recognizer = PatternRecognizer(
        supported_entity="DOCUMENT_ISSUE_DATE",
        name="Document Issue Date Recognizer",
        patterns=[
            Pattern(name="issue_date_ddmmyyyy", regex=r"\b\d{2}[/-]\d{2}[/-]\d{4}\b", score=0.3),
            Pattern(name="issue_date_yyyymmdd", regex=r"\b\d{4}[/-]\d{2}[/-]\d{2}\b", score=0.3)
        ],
        context=["date of issue", "issued on", "issue", "issued"]
    )
    recognizers.append(issue_date_recognizer)

    expiry_date_recognizer = PatternRecognizer(
        supported_entity="DOCUMENT_EXPIRY_DATE",
        name="Document Expiry Date Recognizer",
        patterns=[
            Pattern(name="doc_expiry_ddmmyyyy", regex=r"\b\d{2}[/-]\d{2}[/-]\d{4}\b", score=0.3),
            Pattern(name="doc_expiry_yyyymmdd", regex=r"\b\d{4}[/-]\d{2}[/-]\d{2}\b", score=0.3)
        ],
        context=["date of expiry", "valid until", "valid up to", "expiry", "until"]
    )
    recognizers.append(expiry_date_recognizer)

    gender_recognizer = PatternRecognizer(
        supported_entity="GENDER",
        name="Gender Recognizer",
        patterns=[
            Pattern(name="gender_pattern", regex=r"\b(?:male|female|transgender|other|non-binary)\b", score=0.3)
        ],
        context=["gender", "sex", "m/f", "gender identity"]
    )
    recognizers.append(gender_recognizer)

    # DIGITAL FOOTPRINTS

    ip_recognizer = PatternRecognizer(
        supported_entity="IP_ADDRESS",
        name="IP Address Recognizer",
        patterns=[
            Pattern(name="ipv4_pattern", regex=r"\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b", score=0.6)
        ],
        context=["ip address", "ipv4", "server ip", "host"]
    )
    recognizers.append(ip_recognizer)

    mac_recognizer = PatternRecognizer(
        supported_entity="MAC_ADDRESS",
        name="MAC Address Recognizer",
        patterns=[
            Pattern(name="mac_pattern", regex=r"\b(?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2}\b", score=0.8)
        ],
        context=["mac address", "hardware address", "device mac"]
    )
    recognizers.append(mac_recognizer)

    # SENSITIVE DATA

    medical_record_recognizer = PatternRecognizer(
        supported_entity="MEDICAL_RECORD",
        name="Medical Record Recognizer",
        patterns=[
            Pattern(name="mrn_pattern", regex=r"\b(?:MRN|MR|PRN)[\s:-]?[A-Z0-9]{6,10}\b", score=0.6)
        ],
        context=["medical record", "patient id", "hospital record", "mrn", "patient number"]
    )
    recognizers.append(medical_record_recognizer)

    return recognizers

# ============================================================================
# ANCHOR DETECTION
# ============================================================================

class FuzzyNameAnchorRecognizer(EntityRecognizer):
    """Supplements the exact ANCHOR_NAME recognizer with fuzzy matching against candidate spans."""

    def __init__(self, anchor_name: str):
        super().__init__(
            supported_entities=["ANCHOR_NAME_FUZZY"],
            name="Fuzzy Anchor Name Recognizer",
            supported_language="en",
        )
        self.anchor_name = anchor_name

    def load(self) -> None:
        pass

    def analyze(self, text, entities, nlp_artifacts=None) -> List[RecognizerResult]:
        if entities and "ANCHOR_NAME_FUZZY" not in entities:
            return []

        results = []
        for candidate, start, end in _name_candidate_spans(text):
            if not _passes_fuzzy_guardrails(self.anchor_name, candidate):
                continue
            score = fuzz.token_sort_ratio(self.anchor_name.lower(), candidate.lower())
            if score >= FUZZY_ANCHOR_NAME_THRESHOLD:
                results.append(
                    RecognizerResult(
                        entity_type="ANCHOR_NAME_FUZZY",
                        start=start,
                        end=end,
                        score=0.85,
                    )
                )
        return results

def create_anchor_recognizers(anchors: Dict[str, str]) -> List[PatternRecognizer]:
    """Create custom recognizers for user-provided anchor values"""
    recognizers = []

    for anchor_type, anchor_value in anchors.items():
        if not anchor_value:
            continue

        if anchor_type == "name":
            escaped_name = re.escape(anchor_value)
            name_pattern = r'\b' + escaped_name + r'\b'

            recognizer = PatternRecognizer(
                supported_entity="ANCHOR_NAME",
                name="Anchor Name Recognizer",
                patterns=[
                    Pattern(name="anchor_name_pattern", regex=name_pattern, score=0.95)
                ],
                context=[]
            )
            recognizers.append(recognizer)
            recognizers.append(FuzzyNameAnchorRecognizer(anchor_value))

        elif anchor_type == "email":
            escaped_email = re.escape(anchor_value)
            recognizer = PatternRecognizer(
                supported_entity="ANCHOR_EMAIL",
                name="Anchor Email Recognizer",
                patterns=[
                    Pattern(name="anchor_email_pattern", regex=escaped_email, score=0.95)
                ],
                context=["email", "mail"]
            )
            recognizers.append(recognizer)

        elif anchor_type == "phone":
            clean_phone = re.sub(r'[\s\-\(\)]', '', anchor_value)
            phone_pattern = r'\b' + re.escape(clean_phone) + r'\b'
            recognizer = PatternRecognizer(
                supported_entity="ANCHOR_PHONE",
                name="Anchor Phone Recognizer",
                patterns=[
                    Pattern(name="anchor_phone_pattern", regex=phone_pattern, score=0.95)
                ],
                context=["phone", "mobile", "contact"]
            )
            recognizers.append(recognizer)

        elif anchor_type == "aadhaar":
            clean_aadhaar = re.sub(r'\s', '', anchor_value)
            if len(clean_aadhaar) == 12:
                pattern = r'\b' + clean_aadhaar[0:4] + r'\s?' + clean_aadhaar[4:8] + r'\s?' + clean_aadhaar[8:12] + r'\b'
                recognizer = PatternRecognizer(
                    supported_entity="ANCHOR_AADHAAR",
                    name="Anchor Aadhaar Recognizer",
                    patterns=[
                        Pattern(name="anchor_aadhaar_pattern", regex=pattern, score=0.95)
                    ],
                    context=["aadhaar", "aadhar", "uid"]
                )
                recognizers.append(recognizer)

        elif anchor_type == "pan":
            pan_pattern = r'\b' + re.escape(anchor_value.upper()) + r'\b'
            recognizer = PatternRecognizer(
                supported_entity="ANCHOR_PAN",
                name="Anchor PAN Recognizer",
                patterns=[
                    Pattern(name="anchor_pan_pattern", regex=pan_pattern, score=0.95)
                ],
                context=["pan", "permanent account"]
            )
            recognizers.append(recognizer)

    return recognizers

def check_anchors_in_text(text: str, anchors: Dict[str, str]) -> Set[str]:
    """Check which anchors appear in the text (case-insensitive, flexible matching)"""
    matched_anchors = set()
    text_lower = text.lower()

    for anchor_type, anchor_value in anchors.items():
        if not anchor_value:
            continue

        if anchor_type == "name":
            name_normalized = ' '.join(anchor_value.lower().strip().split())
            if name_normalized in text_lower:
                matched_anchors.add(f"name:{anchor_value}")

        elif anchor_type == "email":
            if anchor_value.lower() in text_lower:
                matched_anchors.add(f"email:{anchor_value}")

        elif anchor_type == "phone":
            clean_phone = re.sub(r'[\s\-\(\)]', '', anchor_value)
            clean_text = re.sub(r'[\s\-\(\)]', '', text_lower)
            if clean_phone.lower() in clean_text:
                matched_anchors.add(f"phone:{anchor_value}")

        elif anchor_type == "aadhaar":
            clean_aadhaar = re.sub(r'\s', '', anchor_value)
            clean_text = re.sub(r'\s', '', text_lower)
            if clean_aadhaar in clean_text:
                matched_anchors.add(f"aadhaar:{anchor_value}")

        elif anchor_type == "pan":
            if anchor_value.upper() in text.upper():
                matched_anchors.add(f"pan:{anchor_value}")

    return matched_anchors

def check_anchors_in_text_for_prefilter(text: str, anchors: Dict[str, str]) -> Set[str]:
    """File-level has_anchors pre-filter: try the cheap exact substring check first, then fall back to fuzzy name matching."""
    matched_anchors = check_anchors_in_text(text, anchors)

    name_value = anchors.get("name")
    if name_value and not any(m.startswith("name:") for m in matched_anchors):
        fuzzy_hit = name_matches_fuzzy(text, name_value)
        if fuzzy_hit:
            matched_anchors.add(f"name_fuzzy:{name_value}~{fuzzy_hit}")

    return matched_anchors

# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def infer_column_context(column_name: str) -> List[str]:
    """Infer semantic context from column name"""
    col = column_name.lower()
    context_keywords = []

    if any(k in col for k in ["name", "student", "employee", "candidate", "person", "patient"]):
        context_keywords.extend(["name", "person"])

    if any(k in col for k in ["phone", "mobile", "contact", "tel"]):
        context_keywords.extend(["phone", "contact"])

    if "email" in col or "mail" in col:
        context_keywords.append("email")

    if any(k in col for k in ["address", "location", "residence", "street"]):
        context_keywords.extend(["address", "location"])

    if any(k in col for k in ["dob", "birth"]):
        context_keywords.extend(["date of birth", "dob"])

    if "gender" in col or "sex" in col:
        context_keywords.append("gender")

    if any(k in col for k in ["aadhaar", "aadhar"]):
        context_keywords.extend(["aadhaar", "uid"])

    if "pan" in col and "company" not in col:
        context_keywords.extend(["pan", "permanent account number"])

    if any(k in col for k in ["passport"]):
        context_keywords.append("passport")

    if any(k in col for k in ["vehicle", "registration", "car", "rc"]):
        context_keywords.extend(["vehicle", "registration"])

    if any(k in col for k in ["voter", "epic"]):
        context_keywords.extend(["voter", "voter id"])

    if any(k in col for k in ["license", "licence", "dl"]):
        context_keywords.extend(["driving license", "dl"])

    if any(k in col for k in ["gstin", "gst"]):
        context_keywords.extend(["gstin", "gst"])

    if any(k in col for k in ["account"]) and any(k in col for k in ["bank", "savings", "current"]):
        context_keywords.extend(["bank account", "account number"])

    if any(k in col for k in ["ifsc"]):
        context_keywords.extend(["ifsc", "bank code"])

    if any(k in col for k in ["upi"]):
        context_keywords.extend(["upi", "payment address"])

    if any(k in col for k in ["medical", "patient", "mrn"]):
        context_keywords.extend(["medical record", "patient id"])

    return context_keywords

def create_semantic_context(column_name: str, value: str) -> str:
    """Create semantic context for better PII detection"""
    context_keywords = infer_column_context(column_name)

    if context_keywords:
        context_str = ", ".join(context_keywords)
        return f"{context_str}: {value}. The column '{column_name}' contains the value {value}."
    else:
        return f"The column '{column_name}' contains the value {value}."

# ----------------------------------------------------------------------------
# Record-boundary construction
# ----------------------------------------------------------------------------

def _group_pdf_page_into_blocks(page, file_path_str: str, page_num: int) -> List[TextRecord]:
    """Word-level bounding boxes -> lines (by vertical position) -> blocks (by vertical gap between lines)."""
    words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
    if not words:
        return []

    lines: List[Dict] = []
    for w in sorted(words, key=lambda w: (w['top'], w['x0'])):
        placed = False
        for line in lines:
            if abs(w['top'] - line['top']) <= 3:
                line['words'].append(w)
                line['top'] = min(line['top'], w['top'])
                line['bottom'] = max(line['bottom'], w['bottom'])
                placed = True
                break
        if not placed:
            lines.append({'top': w['top'], 'bottom': w['bottom'], 'words': [w]})

    if not lines:
        return []

    lines.sort(key=lambda l: l['top'])
    for line in lines:
        line['words'].sort(key=lambda w: w['x0'])

    heights = [l['bottom'] - l['top'] for l in lines if l['bottom'] > l['top']]
    median_height = sorted(heights)[len(heights) // 2] if heights else 10.0
    gap_threshold = median_height * 1.5

    blocks = [[lines[0]]]
    for prev, curr in zip(lines, lines[1:]):
        gap = curr['top'] - prev['bottom']
        if gap > gap_threshold:
            blocks.append([curr])
        else:
            blocks[-1].append(curr)

    records = []
    for block_idx, block_lines in enumerate(blocks, start=1):
        text = "\n".join(" ".join(w['text'] for w in line['words']) for line in block_lines)
        if text.strip():
            records.append(TextRecord(
                record_id=f"{file_path_str}::pdf_block::p{page_num}b{block_idx}",
                text=text,
                boundary_type="pdf_block",
                locator=f"page {page_num}, block {block_idx}",
            ))
    return records

def _group_ocr_data_into_blocks(ocr_data: Dict, file_path_str: str, page_num: int) -> List[TextRecord]:
    """Use Tesseract's own block_num/par_num layout segmentation as the record boundary."""
    groups: Dict[Tuple[int, int], List[str]] = defaultdict(list)
    n = len(ocr_data.get('text', []))
    for i in range(n):
        word = (ocr_data['text'][i] or "").strip()
        if not word:
            continue
        key = (ocr_data['block_num'][i], ocr_data['par_num'][i])
        groups[key].append(word)

    records = []
    for idx, (_, words) in enumerate(sorted(groups.items()), start=1):
        text = " ".join(words)
        if text.strip():
            records.append(TextRecord(
                record_id=f"{file_path_str}::pdf_ocr_block::p{page_num}b{idx}",
                text=text,
                boundary_type="pdf_ocr_block",
                locator=f"page {page_num} (ocr), block {idx}",
            ))
    return records

def _txt_sliding_windows(content: str, file_path_str: str, window: int = TXT_WINDOW_LINES) -> List[TextRecord]:
    """Non-overlapping windows of `window` non-empty lines each."""
    lines = content.splitlines()
    non_empty_idx = [i for i, l in enumerate(lines) if l.strip()]

    records = []
    for start_pos in range(0, len(non_empty_idx), window):
        idxs = non_empty_idx[start_pos:start_pos + window]
        if not idxs:
            continue
        window_text = "\n".join(lines[i] for i in idxs)
        records.append(TextRecord(
            record_id=f"{file_path_str}::line_window::{start_pos}",
            text=window_text,
            boundary_type="line_window",
            locator=f"lines {idxs[0] + 1}-{idxs[-1] + 1}",
        ))
    return records

def extract_text_from_file(file_path: Path, anchors: Dict[str, str]) -> Tuple[bool, List[TextRecord]]:
    """Extract text from file as co-occurrence records and check for anchors."""
    suffix = file_path.suffix.lower()
    file_path_str = str(file_path)
    records: List[TextRecord] = []
    has_any_anchor = False

    try:
        if suffix == ".txt":
            content = file_path.read_text(errors="ignore")
            matched = check_anchors_in_text_for_prefilter(content, anchors)
            if matched:
                has_any_anchor = True
            records = _txt_sliding_windows(content, file_path_str)

        elif suffix == ".pdf":
            pages_without_text = []

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text and len(text.strip()) > 50:
                        matched = check_anchors_in_text_for_prefilter(text, anchors)
                        if matched:
                            has_any_anchor = True
                        records.extend(_group_pdf_page_into_blocks(page, file_path_str, page_num))
                    else:
                        pages_without_text.append(page_num)

            if pages_without_text:
                try:
                    kwargs = {'dpi': 300, 'fmt': 'png'}
                    if POPPLER_PATH:
                        kwargs['poppler_path'] = POPPLER_PATH

                    images = convert_from_path(file_path, **kwargs)
                    custom_config = r'--oem 3 --psm 6'

                    for page_num, img in enumerate(images, start=1):
                        if page_num in pages_without_text:
                            try:
                                ocr_data = pytesseract.image_to_data(
                                    img, lang='eng', config=custom_config, output_type=Output.DICT
                                )
                                page_text = " ".join(w for w in ocr_data.get('text', []) if w and w.strip())
                                if page_text.strip():
                                    matched = check_anchors_in_text_for_prefilter(page_text, anchors)
                                    if matched:
                                        has_any_anchor = True
                                    records.extend(
                                        _group_ocr_data_into_blocks(ocr_data, file_path_str, page_num)
                                    )
                            except Exception:
                                continue
                except Exception:
                    pass

        elif suffix == ".docx":
            doc = Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            matched = check_anchors_in_text_for_prefilter(full_text, anchors)
            if matched:
                has_any_anchor = True

            for para_idx, para in enumerate(doc.paragraphs, start=1):
                if para.text.strip():
                    records.append(TextRecord(
                        record_id=f"{file_path_str}::paragraph::{para_idx}",
                        text=para.text,
                        boundary_type="paragraph",
                        locator=f"paragraph {para_idx}",
                    ))

        elif suffix.startswith(".xls"):
            try:
                sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)
            except Exception:
                return False, []

            for sheet_name, df in sheets.items():
                df = df.fillna("")

                sheet_text = df.to_string()
                matched = check_anchors_in_text_for_prefilter(sheet_text, anchors)
                if matched:
                    has_any_anchor = True

                for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
                    row_sentences = []
                    for col in df.columns:
                        value = str(row[col]).strip()
                        if value:
                            row_sentences.append(create_semantic_context(col, value))

                    if row_sentences:
                        row_text = " ".join(row_sentences)
                        records.append(TextRecord(
                            record_id=f"{file_path_str}::row::{sheet_name}:{row_idx}",
                            text=row_text,
                            boundary_type="row",
                            locator=f"sheet '{sheet_name}', row {row_idx}",
                        ))

        elif suffix == ".csv":
            full_df = pd.read_csv(file_path, dtype=str)
            full_text = full_df.to_string()

            matched = check_anchors_in_text_for_prefilter(full_text, anchors)
            if matched:
                has_any_anchor = True

            row_idx = 0
            for chunk_df in pd.read_csv(file_path, chunksize=500, dtype=str):
                chunk_df = chunk_df.fillna("")
                for _, row in chunk_df.iterrows():
                    row_idx += 1
                    row_sentences = []
                    for col in chunk_df.columns:
                        value = str(row[col]).strip()
                        if value:
                            row_sentences.append(create_semantic_context(col, value))

                    if row_sentences:
                        row_text = " ".join(row_sentences)
                        records.append(TextRecord(
                            record_id=f"{file_path_str}::row::{row_idx}",
                            text=row_text,
                            boundary_type="row",
                            locator=f"row {row_idx}",
                        ))

    except Exception:
        return False, []

    return has_any_anchor, records

def extract_ml_entities(text):
    if nlp is None:
        return []

    doc = nlp(text)
    return [
        PIIInstance(
            entity_type="ML_PERSON",
            value=ent.text,
            file_path="",
            score=0.9
        )
        for ent in doc.ents if ent.label_ == "PERSON"
    ]

def extract_label_matched_names(text: str, file_path: str, record_id: str) -> List[PIIInstance]:
    """Field-label ('Name:', 'Account Holder:', ...) and Indian relational-marker (S/o, D/o, W/o, C/o) name extraction."""
    instances = []
    seen = set()
    for span_text, start, end in _find_label_matched_name_spans(text):
        key = span_text.lower()
        if key in seen:
            continue
        seen.add(key)

        context_start = max(0, start - 50)
        context_end = min(len(text), end + 50)
        context = text[context_start:context_end].replace('\n', ' ').strip()

        instances.append(PIIInstance(
            entity_type="LABEL_MATCHED_NAME",
            value=span_text,
            file_path=file_path,
            context=context,
            score=0.95,
            record_id=record_id,
        ))
    return instances

# ============================================================================
# PII ANALYSIS
# ============================================================================

def analyze_text_for_pii(
    text: str,
    analyzer: AnalyzerEngine,
    file_path: str,
    matched_anchors: Set[str]
) -> Tuple[List[PIIInstance], List[PIIInstance]]:
    """Hybrid PII detection: Presidio (rule-based + pretrained NLP) plus spaCy NER as an internal candidate source.

    Returns (output_instances, ml_person_candidates).
    """

    instances = []
    ml_person_candidates = []

    try:
        # =========================
        # 1. PRESIDIO DETECTION
        # =========================
        results = analyzer.analyze(
            text=text,
            language="en",
            score_threshold=SCORE_THRESHOLD
        )

        for result in results:
            entity_type = result.entity_type

            if entity_type not in ALLOWED_ENTITIES and not entity_type.startswith("ANCHOR_"):
                continue

            value = text[result.start:result.end]

            context_start = max(0, result.start - 50)
            context_end = min(len(text), result.end + 50)
            context = text[context_start:context_end].replace('\n', ' ').strip()

            instance = PIIInstance(
                entity_type=entity_type,
                value=value,
                file_path=file_path,
                context=context,
                score=result.score
            )

            instances.append(instance)

        # =========================
        # 2. ML NER (spaCy) candidates - internal only
        # =========================
        doc = _get_spacy_doc(text)
        if doc is not None:
            for ent in doc.ents:
                if ent.label_ == "PERSON":

                    value = ent.text.strip()

                    if any(i.value.lower() == value.lower() for i in instances):
                        continue

                    candidate = PIIInstance(
                        entity_type="ML_PERSON",
                        value=value,
                        file_path=file_path,
                        context=text[max(0, ent.start_char-50):ent.end_char+50],
                        score=0.9
                    )

                    ml_person_candidates.append(candidate)

        # =========================
        # 3. OPTIONAL: BOOST SCORE IF ANCHOR PRESENT
        # =========================
        if matched_anchors:
            for inst in instances:
                inst.score = min(inst.score + 0.1, 1.0)

    except Exception:
        pass

    return instances, ml_person_candidates

# ============================================================================
# CLUSTERING & CONFIDENCE SCORING
# ============================================================================

def calculate_cluster_confidence(cluster: PIICluster, total_files_scanned: int) -> float:
    """Calculate confidence score for a PII cluster."""
    confidence = 0.0

    # Factor 1: Anchor matches (0-40 points)
    num_anchors = len(cluster.anchor_matches)
    if num_anchors >= 3:
        confidence += 40
    elif num_anchors == 2:
        confidence += 30
    elif num_anchors == 1:
        confidence += 20

    # Factor 2: Strong identifiers (0-30 points)
    strong_identifiers = {
        'IN_AADHAAR', 'IN_PAN', 'IN_PASSPORT', 'IN_VOTER_ID',
        'IN_DRIVING_LICENSE', 'ANCHOR_AADHAAR', 'ANCHOR_PAN'
    }
    unique_pii = cluster.get_unique_pii()
    strong_count = sum(1 for entity_type in unique_pii.keys() if entity_type in strong_identifiers)
    confidence += min(strong_count * 10, 30)

    # Factor 3: Diversity of PII types (0-20 points)
    num_pii_types = len(unique_pii)
    confidence += min(num_pii_types * 3, 20)

    # Factor 4: Multiple files (0-10 points)
    num_files = len(cluster.files)
    if num_files > 1:
        confidence += min(num_files * 3, 10)

    # Normalize to 0-1 scale
    return min(confidence / 100.0, 1.0)

def create_pii_cluster(pii_instances: List[PIIInstance], matched_anchors: Set[str],
                       total_files: int) -> PIICluster:
    """Create a single PII cluster from instances."""
    cluster = PIICluster()
    cluster.anchor_matches = matched_anchors

    for instance in pii_instances:
        cluster.add_instance(instance)

    cluster.confidence = calculate_cluster_confidence(cluster, total_files)

    return cluster

# ============================================================================
# ANCHOR PROXIMITY GRAPH
# ============================================================================

# Free-text / NLP-derived entity types are never used as cross-record bridges;
# structured identifiers (phone, Aadhaar, email, ANCHOR_*, ...) keep the
# cross-record merge since a repeated value there is real signal.
_RECORD_SCOPED_ENTITY_TYPES = {"LABEL_MATCHED_NAME"}

def _node_key(instance: PIIInstance) -> Tuple[str, ...]:
    base = (instance.entity_type, instance.value.lower().strip())
    if instance.entity_type in _RECORD_SCOPED_ENTITY_TYPES:
        return base + (instance.record_id,)
    return base

def _build_file_graph(instances: List[PIIInstance]) -> "nx.Graph":
    """Nodes = distinct (entity_type, normalized value); edges = node pairs that co-occur in the same record."""
    graph = nx.Graph()
    by_record: Dict[str, List[PIIInstance]] = defaultdict(list)
    for inst in instances:
        by_record[inst.record_id].append(inst)
        graph.add_node(_node_key(inst))

    for record_instances in by_record.values():
        keys = list({_node_key(i) for i in record_instances})
        for i in range(len(keys)):
            for j in range(i + 1, len(keys)):
                graph.add_edge(keys[i], keys[j])

    return graph

def _filter_to_anchor_components(instances: List[PIIInstance]) -> List[PIIInstance]:
    """Keep only instances whose node lies in a connected component that contains at least one ANCHOR_* node."""
    if not instances:
        return []

    graph = _build_file_graph(instances)
    anchor_keys = {_node_key(i) for i in instances if i.entity_type.startswith("ANCHOR_")}
    if not anchor_keys:
        return []

    retained_keys: Set[Tuple[str, str]] = set()
    for component in nx.connected_components(graph):
        if component & anchor_keys:
            retained_keys.update(component)

    return [i for i in instances if _node_key(i) in retained_keys]

# ============================================================================
# CROSS-DOCUMENT NAME RECONCILIATION
# ============================================================================

_NAME_SOURCE_TYPES = ("ANCHOR_NAME", "LABEL_MATCHED_NAME", "ANCHOR_NAME_FUZZY", "ML_PERSON")
_NAME_SOURCE_PRIORITY = {t: i for i, t in enumerate(_NAME_SOURCE_TYPES)}

def _reconcile_names(instances: List[PIIInstance]) -> List[Dict]:
    """Group name-variant candidates across the whole scan into one canonical name plus known variants."""
    candidates = [i for i in instances if i.entity_type in _NAME_SOURCE_TYPES]
    if not candidates:
        return []

    by_value: Dict[str, Dict] = {}
    for inst in candidates:
        key = inst.value.lower().strip()
        if key not in by_value:
            by_value[key] = {"value": inst.value, "entity_type": inst.entity_type}
        elif _NAME_SOURCE_PRIORITY[inst.entity_type] < _NAME_SOURCE_PRIORITY[by_value[key]["entity_type"]]:
            by_value[key]["entity_type"] = inst.entity_type

    values = list(by_value.keys())
    graph = nx.Graph()
    graph.add_nodes_from(values)
    for i in range(len(values)):
        for j in range(i + 1, len(values)):
            score = fuzz.token_sort_ratio(values[i], values[j])
            if score >= FUZZY_RECONCILIATION_THRESHOLD:
                graph.add_edge(values[i], values[j])

    results = []
    for component in nx.connected_components(graph):
        members = [by_value[v] for v in component]
        members.sort(key=lambda m: (
            _NAME_SOURCE_PRIORITY[m["entity_type"]],
            -len(m["value"].split()),
            -len(m["value"]),
        ))
        canonical = members[0]["value"]
        variants = sorted({m["value"] for m in members} - {canonical})
        sources = sorted({m["entity_type"] for m in members})
        results.append({
            "canonical": canonical,
            "variants": variants,
            "sources": sources,
        })

    return results

# ============================================================================
# MAIN DISCOVERY ENGINE
# ============================================================================

def discover_pii(folder_path: str, **anchors) -> Dict[str, any]:
    """Main PII discovery function with clustering."""

    if not anchors:
        raise ValueError("At least one anchor must be provided (name, email, phone, aadhaar, or pan)")

    root = Path(folder_path)
    if not root.exists():
        raise ValueError(f"Folder path does not exist: {folder_path}")

    analyzer = AnalyzerEngine()

    comprehensive_recognizers = create_comprehensive_recognizers()
    for recognizer in comprehensive_recognizers:
        analyzer.registry.add_recognizer(recognizer)

    anchor_recognizers = create_anchor_recognizers(anchors)
    for recognizer in anchor_recognizers:
        analyzer.registry.add_recognizer(recognizer)

    print("\n" + "="*80)
    print("PII DISCOVERY")
    print("="*80)
    print(f"\nScanning: {folder_path}\n")

    all_pii_instances = []
    all_ml_person_candidates: List[PIIInstance] = []
    files_scanned = 0
    files_with_anchors = 0
    files_with_pii = 0
    global_anchor_matches = set()

    print("Scanning files...")

    for file_path in root.rglob("*"):
        if not file_path.is_file():
            continue

        files_scanned += 1

        has_anchors, records = extract_text_from_file(file_path, anchors)

        if not has_anchors:
            continue

        files_with_anchors += 1

        raw_instances: List[PIIInstance] = []
        ml_person_candidates_this_file: List[PIIInstance] = []
        for record in records:
            if not record.text.strip():
                continue

            chunk_anchors = check_anchors_in_text(record.text, anchors)
            global_anchor_matches.update(chunk_anchors)

            record_instances, record_ml_candidates = analyze_text_for_pii(
                record.text,
                analyzer,
                str(file_path),
                chunk_anchors
            )
            for inst in record_instances:
                inst.record_id = record.record_id
            for cand in record_ml_candidates:
                cand.record_id = record.record_id
            raw_instances.extend(record_instances)
            ml_person_candidates_this_file.extend(record_ml_candidates)

        filtered_instances = _filter_to_anchor_components(raw_instances)

        retained_record_ids = {inst.record_id for inst in filtered_instances}
        records_by_id = {r.record_id: r for r in records}
        for record_id in retained_record_ids:
            record = records_by_id.get(record_id)
            if record is None:
                continue
            filtered_instances.extend(
                extract_label_matched_names(record.text, str(file_path), record.record_id)
            )

        all_ml_person_candidates.extend(
            cand for cand in ml_person_candidates_this_file
            if cand.record_id in retained_record_ids
        )

        if filtered_instances:
            files_with_pii += 1
            all_pii_instances.extend(filtered_instances)

    print(f"Scanned {files_scanned} files.\n")

    cluster = create_pii_cluster(all_pii_instances, global_anchor_matches, files_scanned)

    results = {
        'search_timestamp': datetime.now().isoformat(),
        'search_parameters': {k: v for k, v in anchors.items() if v},
        'statistics': {
            'total_files_scanned': files_scanned,
            'files_with_anchors': files_with_anchors,
            'files_with_pii': files_with_pii,
            'total_pii_instances': len(all_pii_instances),
            'unique_files_with_pii': len(cluster.files)
        },
        'cluster': {
            'confidence': cluster.confidence,
            'confidence_level': (
                'HIGH' if cluster.confidence >= CONFIDENCE_HIGH else
                'MEDIUM' if cluster.confidence >= CONFIDENCE_MEDIUM else
                'LOW'
            ),
            'matched_anchors': list(cluster.anchor_matches),
            'pii_summary': cluster.get_pii_summary(),
            'unique_pii_values': {k: list(v) for k, v in cluster.get_unique_pii().items()},
            'files': list(cluster.files)
        },
        'identity_resolution': {
            'canonical_names': _reconcile_names(all_pii_instances + all_ml_person_candidates)
        },
        'detailed_instances': []
    }

    instances_by_file = defaultdict(list)
    for instance in all_pii_instances:
        instances_by_file[instance.file_path].append(instance)

    for file_path, instances in instances_by_file.items():
        file_data = {
            'file': file_path,
            'pii_count': len(instances),
            'pii_types': defaultdict(list)
        }

        for instance in instances:
            file_data['pii_types'][instance.entity_type].append({
                'value': instance.value,
                'context': instance.context,
                'score': instance.score
            })

        file_data['pii_types'] = dict(file_data['pii_types'])
        results['detailed_instances'].append(file_data)

    return results

# ============================================================================
# REPORTING
# ============================================================================

def print_discovery_report(results: Dict):
    """Print a simplified discovery report - only file paths and PII types"""

    print("\n" + "="*80)
    print("PII DISCOVERY REPORT")
    print("="*80 + "\n")

    detailed_instances = results.get('detailed_instances', [])

    if not detailed_instances:
        print("No PII found for the specified person.")
        print("\n" + "="*80 + "\n")
        return

    print(f"Found PII in {len(detailed_instances)} file(s):\n")

    # Print each file with its PII types
    for file_data in detailed_instances:
        file_path = file_data['file']
        pii_types = file_data['pii_types']

        print(f"FILE: {file_path}")

        for pii_type in sorted(pii_types.keys()):
            print(f"  - {pii_type}")

        print()

    print("="*80 + "\n")

def save_report_json(results: Dict, output_path: str):
    """Save detailed report as JSON"""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    print(f"Detailed report saved to: {output_path}")
